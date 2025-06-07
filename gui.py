import sys
import os
# from PyQt5 import QtWidgets, uic # PyQt5 관련 임포트 주석 처리 또는 삭제
# from PyQt5.QtWidgets import QFileDialog, QMessageBox
# from PyQt5.QtCore import QThread, pyqtSignal
import subprocess
import json
from pathlib import Path
import logging

# PyQt6 임포트
from PyQt6 import QtWidgets, QtCore
from PyQt6.QtWidgets import QFileDialog, QMessageBox
from PyQt6.QtCore import QThread, pyqtSignal

# pyuic6로 변환된 UI 파일에서 클래스 임포트
# pyuic6 실행 시 -o 옵션으로 지정한 파일 이름과 Ui_MainWindow 클래스 이름이 맞는지 확인하세요.
from gui_elements import Ui_MainWindow

SETTINGS_PATH = str(Path.home() / '.capture_gui_settings.json')
logger = logging.getLogger('main_gui')
if not logger.hasHandlers():
    ch = logging.StreamHandler()
    formatter = logging.Formatter('%(levelname)s: %(message)s')
    ch.setFormatter(formatter)
    logger.addHandler(ch)
logger.setLevel(logging.INFO)

# 작업 스레드 클래스 - 캡처, PDF 변환, OCR 작업을 비동기로 처리
# PyQt6의 QThread를 상속받아 구현
# 작업 진행 상황과 결과를 메인 GUI에 전달하기 위한 시그널 정의
class Worker(QThread):
    # PyQt6 시그널 선언 방식 (PyQt5와 동일)
    log_signal = pyqtSignal(str)
    finished_signal = pyqtSignal()
    error_signal = pyqtSignal(str)

    def __init__(self, params):
        """
        초기화 메서드
        
        Args:
            params (dict): 작업에 필요한 모든 매개변수를 포함하는 딕셔너리
        """
        super().__init__()
        self.params = params
        self.processes = []  # Track subprocesses for cleanup
        self._is_paused = False
        self._is_running = True
        self._lock = QtCore.QMutex()

    def pause(self):
        """Pause the worker"""
        self._lock.lock()
        self._is_paused = True
        self._lock.unlock()
        self.log_signal.emit('작업이 일시중지되었습니다.')
    
    def resume(self):
        """Resume the worker"""
        self._lock.lock()
        self._is_paused = False
        self._lock.unlock()
        self.log_signal.emit('작업이 재개되었습니다.')
    
    def stop(self):
        """Stop the worker"""
        self._lock.lock()
        self._is_running = False
        self._is_paused = False  # Ensure we're not stuck in a paused state
        self._lock.unlock()
        self.log_signal.emit('작업을 중지하고 있습니다...')
        self.cleanup_processes()
    
    def _check_pause(self):
        """Check if paused and wait if needed"""
        while True:
            self._lock.lock()
            if not self._is_paused:
                self._lock.unlock()
                return True
            self._lock.unlock()
            self.msleep(100)  # Sleep for 100ms
            
            # Check if we should stop while paused
            self._lock.lock()
            if not self._is_running:
                self._lock.unlock()
                return False
            self._lock.unlock()
    
    def run(self):
        """
        비동기 작업 실행 메서드
        캡처 -> PDF 변환 -> OCR의 순서로 작업을 실행
        각 단계별로 진행 상황을 메인 GUI에 보고
        """
        try:
            # Check if we should run
            self._lock.lock()
            if not self._is_running:
                self._lock.unlock()
                return
            self._lock.unlock()
            
            # 단계별 실행
            if self.params['capture']:
                # Check if paused or stopped
                if not self._check_pause():
                    return
                    
                """1단계: 화면 캡처 작업 실행"""
                self.log_signal.emit('[1/3] 캡처 시작...')
                # Build the command with updated parameter names
                cmd = [
                    'python3', 'shot.py',
                    '--app', self.params['app_name'],
                    '--label', self.params['window_label'],
                    '--output-dir', self.params['output_dir'],
                    '--book', self.params['book'],
                    '--start-page', str(self.params['start_page']),
                    '--end-page', str(self.params['end_page']),
                    '--next', self.params['next_action'],
                    '--delay', str(self.params['delay'])
                ]
                
                # Add optional parameters if they have non-zero values
                if self.params['width'] > 0:
                    cmd.extend(['--width', str(self.params['width'])])
                if self.params['height'] > 0:
                    cmd.extend(['--height', str(self.params['height'])])
                    
                # Add margin parameters if they have non-zero values
                if self.params['top_margin'] > 0:
                    cmd.extend(['--top-margin', str(self.params['top_margin'])])
                if self.params['bottom_margin'] > 0:
                    cmd.extend(['--bottom-margin', str(self.params['bottom_margin'])])
                if self.params['left_margin'] > 0:
                    cmd.extend(['--left-margin', str(self.params['left_margin'])])
                if self.params['right_margin'] > 0:
                    cmd.extend(['--right-margin', str(self.params['right_margin'])])
                process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
                self.processes.append(process)
                stdout, stderr = process.communicate()
                result = subprocess.CompletedProcess(cmd, process.returncode, stdout, stderr)
                self.log_signal.emit(result.stdout)
                if result.returncode != 0:
                    self.error_signal.emit(result.stderr)
                    return
            if self.params['pdf']:
                self.log_signal.emit('[2/3] PDF 변환 시작...')
                input_dir = os.path.join(self.params['output_dir'], self.params['book']) if self.params['book'] else self.params['output_dir']
                cmd = [
                    'python3', 'img_to_pdf.py',
                    '-i', input_dir,
                    '-o', os.path.join(self.params['output_dir'], f"{self.params['book']}.pdf") if self.params['book'] else os.path.join(self.params['output_dir'], 'output.pdf')
                ]
                if self.params['pdf_merge']:
                    cmd.append('--merge')
                process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
                self.processes.append(process)
                stdout, stderr = process.communicate()
                result = subprocess.CompletedProcess(cmd, process.returncode, stdout, stderr)
                self.log_signal.emit(result.stdout)
                if result.returncode != 0:
                    self.error_signal.emit(result.stderr)
                    return
            if self.params['ocr']:
                self.log_signal.emit('[3/4] OCR 시작...')
                input_dir = os.path.join(self.params['output_dir'], self.params['book']) if self.params['book'] else self.params['output_dir']
                output_txt = os.path.join(self.params['output_dir'], f"{self.params['book']}.txt") if self.params['book'] else os.path.join(self.params['output_dir'], 'output.txt')
                cmd = [
                    'python3', 'input_to_txt.py',
                    '-i', input_dir,
                    '-o', output_txt
                ]
                if self.params['ocr_merge']:
                    cmd.append('--merge')
                process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
                self.processes.append(process)
                stdout, stderr = process.communicate()
                result = subprocess.CompletedProcess(cmd, process.returncode, stdout, stderr)
                self.log_signal.emit(result.stdout)
                if result.returncode != 0:
                    self.error_signal.emit(result.stderr)
                    return
                    
                # DOCX 변환 단계 추가
                if self.params.get('docx', False) and os.path.exists(output_txt):
                    self.log_signal.emit('[4/4] DOCX 변환 시작...')
                    try:
                        import docx
                        from docx.shared import Pt
                        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                        
                        # Create output path with .docx extension
                        docx_file = os.path.splitext(output_txt)[0] + '.docx'
                        
                        # Create a new Document
                        doc = docx.Document()
                        
                        # Set default font
                        style = doc.styles['Normal']
                        font = style.font
                        font.name = 'Malgun Gothic'  # Use Korean font
                        font.size = Pt(10.5)
                        
                        # Read the text file
                        with open(output_txt, 'r', encoding='utf-8') as f:
                            lines = f.readlines()
                        
                        # Add each line to the document
                        for line in lines:
                            # Skip empty lines
                            if not line.strip():
                                continue
                                
                            # Add paragraph with proper alignment
                            p = doc.add_paragraph()
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.line_spacing = 1.5
                            
                            # Add text with formatting
                            run = p.add_run(line.strip())
                            run.font.name = 'Malgun Gothic'
                            run.font.size = Pt(10.5)
                        
                        # Save the document
                        doc.save(docx_file)
                        self.log_signal.emit(f'DOCX 파일이 생성되었습니다: {docx_file}')
                        
                    except Exception as e:
                        self.log_signal.emit(f'DOCX 변환 중 오류가 발생했습니다: {str(e)}')
                        return
                        
            self.finished_signal.emit()
        except Exception as e:
            self.error_signal.emit(str(e))
    
    def cleanup_processes(self):
        """Clean up any running subprocesses"""
        for process in self.processes:
            try:
                if process.poll() is None:  # Check if process is still running
                    process.terminate()  # Try to terminate gracefully first
                    try:
                        process.wait(timeout=2)  # Give it a moment to terminate
                    except subprocess.TimeoutExpired:
                        process.kill()  # Force kill if it doesn't terminate
            except Exception as e:
                logger.error(f"Error cleaning up process: {e}")


# 메인 GUI 클래스 - PyQt6를 사용한 메인 윈도우 구현
class DocumentProcessorGUI(QtWidgets.QMainWindow):
    def __init__(self):
        """메인 윈도우 초기화 메서드
        UI 설정, 이벤트 핸들러 연결, 기본 설정 로드 등을 수행
        """
        super().__init__()
        # uic.loadUi('gui.ui', self) # PyQt5 UI 로딩 방식 주석 처리 또는 삭제

        # PyQt6에서 pyuic6로 변환된 클래스를 사용하여 UI 로드
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        
        # 윈도우 종료 이벤트 연결
        self.destroyed.connect(self.cleanup)

        # 이제 self.ui.위젯이름 형태로 위젯에 접근합니다.
        self.ui.browseOutputDirButton.clicked.connect(self.browse_output_dir)
        self.ui.runButton.clicked.connect(self.start_workflow)
        
        # Add control buttons
        self.pauseButton = QtWidgets.QPushButton("일시정지")
        self.pauseButton.setObjectName("pauseButton")
        self.pauseButton.setEnabled(False)
        self.pauseButton.clicked.connect(self.toggle_pause)
        
        self.stopButton = QtWidgets.QPushButton("중지")
        self.stopButton.setObjectName("stopButton")
        self.stopButton.setEnabled(False)
        self.stopButton.clicked.connect(self.stop_workflow)
        
        # Find the layout containing the runButton and add new buttons there
        # Assuming runButton is in a vertical layout directly under centralwidget or similar
        # This part might need adjustment based on actual UI structure in ui_main_gui.py
        # For now, let's assume it's in a simple layout that we can modify
        # A more robust solution would involve modifying the .ui file and recompiling
        
        # Get the current layout of the centralwidget
        main_layout = self.ui.centralwidget.layout()
        if main_layout is None:
            main_layout = QtWidgets.QVBoxLayout(self.ui.centralwidget)
            self.ui.centralwidget.setLayout(main_layout)
            
        # Find the index of runButton in the layout
        run_button_index = -1
        for i in range(main_layout.count()):
            item = main_layout.itemAt(i)
            if item.widget() == self.ui.runButton:
                run_button_index = i
                break
        
        if run_button_index != -1:
            # Remove the runButton from its current position
            main_layout.removeWidget(self.ui.runButton)
            
            # Create a new horizontal layout for control buttons
            control_button_layout = QtWidgets.QHBoxLayout()
            control_button_layout.addWidget(self.ui.runButton)
            control_button_layout.addWidget(self.pauseButton)
            control_button_layout.addWidget(self.stopButton)
            
            # Insert the new horizontal layout at the runButton's original position
            main_layout.insertLayout(run_button_index, control_button_layout)
        else:
            # Fallback if runButton is not found in the main layout (e.g., it's in a sub-groupbox)
            # This is a simpler approach, assuming runButton is directly in a layout that can be appended to
            # If the UI is complex, manual .ui file editing is better.
            button_layout = QtWidgets.QHBoxLayout()
            button_layout.addWidget(self.ui.runButton)
            button_layout.addWidget(self.pauseButton)
            button_layout.addWidget(self.stopButton)
            main_layout.addLayout(button_layout)

        # 메뉴 액션 연결 (self.ui.action이름 형태 사용)
        self.ui.actionSave_Settings.triggered.connect(self.save_settings)
        self.ui.actionLoad_Settings.triggered.connect(self.load_settings)
        
        # Track if workflow is running/paused
        self.workflow_running = False
        self.workflow_paused = False

        # 로그 영역 (self.ui.logTextEdit 형태 사용)
        self.ui.logTextEdit.setReadOnly(True)

        self.worker = None

        # defaults 리스트에서도 위젯 접근 방식을 self.ui.위젯 형태로 변경
        self.defaults = [
            (self.ui.appNameLineEdit, 'Windows App', 'setText', 'text'),
            (self.ui.windowLabelLineEdit, 'Mini PC', 'setText', 'text'),
            (self.ui.outputDirLineEdit, str(Path.home() / 'Documents'), 'setText', 'text'),
            (self.ui.bookLineEdit, 'bookTitle', 'setText', 'text'),
            (self.ui.startSpinBox, 1, 'setValue', 'value'),
            (self.ui.noSpinBox, 5, 'setValue', 'value'),  # Using noSpinBox to match UI definition
            (self.ui.delaySpinBox, 0.1, 'setValue', 'value'),
            (self.ui.widthSpinBox, 2880, 'setValue', 'value'),
            (self.ui.heightSpinBox, 1800, 'setValue', 'value'),
            (self.ui.topSpinBox, 50, 'setValue', 'value'),
            (self.ui.bottomSpinBox, 55, 'setValue', 'value'),
            (self.ui.leftSpinBox, 0, 'setValue', 'value'),
            (self.ui.rightSpinBox, 0, 'setValue', 'value'),
            (self.ui.tessPathLineEdit, '', 'setText', 'text'),
            (self.ui.langLineEdit, 'kor+eng', 'setText', 'text'),
            (self.ui.ocrCheckBox, True, 'setChecked', 'checked'),
            (self.ui.pdfCheckBox, True, 'setChecked', 'checked'),
            (self.ui.textCheckBox, True, 'setChecked', 'checked'),
            (self.ui.mergeCheckBox, True, 'setChecked', 'checked')
        ]

        self.load_settings()

        # 위젯 접근 방식 변경 (self.ui.위젯 형태 사용)
        self.ui.widthSpinBox.setEnabled(True)
        self.ui.heightSpinBox.setEnabled(True)
        self.ui.topSpinBox.setEnabled(True)
        self.ui.bottomSpinBox.setEnabled(True)
        self.ui.leftSpinBox.setEnabled(True)
        self.ui.rightSpinBox.setEnabled(True)

    def load_settings(self):
        """
        설정 파일에서 이전에 저장된 설정값을 불러오는 메서드
        
        설정 파일이 없거나 오류가 발생한 경우 기본값을 사용합니다.
        모든 위젯의 상태를 설정 파일의 값으로 업데이트합니다.
        """
        logger.info(f"[LOAD] Loading settings from {SETTINGS_PATH}")
        try:
            with open(SETTINGS_PATH, 'r', encoding='utf-8') as f:
                data = json.load(f)
            logger.info(f"[LOAD] File content: {data}")
            for widget, default, set_method, value_type in self.defaults:
                key = widget.objectName()
                if key not in data:
                    logger.warning(f"[LOAD] {key} not found in settings, using default: {default}")
                    getattr(widget, set_method)(default)
                    continue
                value = data[key]
                try:
                    if value_type == 'checked':
                        getattr(widget, set_method)(bool(value))
                    else:
                        getattr(widget, set_method)(value)
                    logger.info(f"[LOAD] Set {key} to {value}")
                except Exception as e:
                    logger.error(f"[LOAD] Failed to set {key}: {e}")
        except Exception as e:
            logger.warning(f"[LOAD] Failed to load settings: {e}. Applying defaults.")
            for widget, default, set_method, _ in self.defaults:
                try:
                    getattr(widget, set_method)(default)
                    logger.info(f"[LOAD] Set {widget.objectName()} to default {default}")
                except Exception as e2:
                    logger.error(f"[LOAD] Failed to set default for {widget.objectName()}: {e2}")
        # After loading, print all widget values
        for widget, _, _, value_type in self.defaults:
            try:
                key = widget.objectName()
                if value_type == 'text':
                    logger.info(f"[LOAD] {key} value: {widget.text()}")
                elif value_type == 'value':
                    logger.info(f"[LOAD] {key} value: {widget.value()}")
                elif value_type == 'checked':
                    logger.info(f"[LOAD] {key} value: {widget.isChecked()}")
            except Exception as e:
                logger.error(f"[LOAD] Failed to read {key}: {e}")

    def save_settings(self):
        """
        현재 위젯들의 상태를 설정 파일에 저장하는 메서드
        
        모든 위젯의 현재 값을 JSON 형식으로 저장합니다.
        저장 위치는 SETTINGS_PATH에 정의된 경로입니다.
        """
        data = {}  # 설정값을 저장할 딕셔너리 초기화
        for widget, _, _, value_type in self.defaults:
            key = widget.objectName()
            try:
                if value_type == 'text':
                    data[key] = widget.text()
                elif value_type == 'value':
                    data[key] = widget.value()
                elif value_type == 'checked':
                    data[key] = widget.isChecked()
                logger.info(f"[SAVE] {key} value to save: {data[key]}")
            except Exception as e:
                logger.error(f"[SAVE] Failed to get value for {key}: {e}")
        try:
            with open(SETTINGS_PATH, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            logger.info(f"[SAVE] Saved settings to {SETTINGS_PATH}")
            logger.info(f"[SAVE] File content: {data}")
        except Exception as e:
            logger.error(f"[SAVE] Failed to save settings: {e}")

    # closeEvent 메소드는 변경 필요 없음
    def closeEvent(self, event):
        """
        윈도우 종료 이벤트 핸들러
        
        Args:
            event: 종료 이벤트 객체
        """
        self.save_settings()  # 종료 전 설정 저장
        super().closeEvent(event)

    def browse_output_dir(self):
        """
        출력 디렉토리를 선택하기 위한 다이얼로그를 여는 메서드
        
        사용자가 선택한 디렉토리 경로를 출력 디렉토리 입력 필드에 설정합니다.
        """
        # QFileDialog는 PyQt6.QtWidgets에서 임포트하므로 변경 없음
        dir_path = QFileDialog.getExistingDirectory(self, 'Select Output Directory')
        if dir_path:
            self.ui.outputDirLineEdit.setText(dir_path) # self.ui 추가

    def get_params(self):
        # 위젯 접근 방식을 self.ui.위젯 형태로 변경
        output_dir = self.ui.outputDirLineEdit.text().strip()
        start_page = self.ui.startSpinBox.value()
        end_page = self.ui.noSpinBox.value()
        
        # Ensure end_page is not less than start_page
        if end_page < start_page:
            end_page = start_page
            self.ui.noSpinBox.setValue(end_page)
            
        params = {
            'app_name': self.ui.appNameLineEdit.text().strip(),
            'window_label': self.ui.windowLabelLineEdit.text().strip(),
            'output_dir': output_dir,
            'book': self.ui.bookLineEdit.text().strip(),
            'start_page': start_page,
            'end_page': end_page,
            'next_action': 'right',  # Default next action
            'delay': self.ui.delaySpinBox.value(),
            'width': self.ui.widthSpinBox.value(),
            'height': self.ui.heightSpinBox.value(),
            'top_margin': self.ui.topSpinBox.value(),
            'bottom_margin': self.ui.bottomSpinBox.value(),
            'left_margin': self.ui.leftSpinBox.value(),
            'right_margin': self.ui.rightSpinBox.value(),
            'tess_path': self.ui.tessPathLineEdit.text().strip(),
            'lang': self.ui.langLineEdit.text().strip(),
            'capture': True,
            'pdf': self.ui.pdfCheckBox.isChecked(),
            'pdf_merge': self.ui.mergeCheckBox.isChecked(),
            'ocr': self.ui.ocrCheckBox.isChecked(),
            'ocr_merge': self.ui.mergeCheckBox.isChecked(),
            'docx': self.ui.docxCheckBox.isChecked(),  # DOCX conversion flag
        }
        return params

    def cleanup(self):
        """Clean up resources when the application is closing"""
        if hasattr(self, 'worker') and self.worker:
            if self.worker.isRunning():
                self.worker.cleanup_processes()
                self.worker.quit()
                self.worker.wait(2000)  # Wait up to 2 seconds for clean exit
                if self.worker.isRunning():
                    self.worker.terminate()
                    self.worker.wait()
    
    def toggle_pause(self):
        """Toggle pause/resume of the workflow"""
        if not hasattr(self, 'worker') or not self.worker:
            return
            
        if self.workflow_paused:
            self.worker.resume()
            self.pauseButton.setText("일시정지")
            self.workflow_paused = False
        else:
            self.worker.pause()
            self.pauseButton.setText("재개")
            self.workflow_paused = True
    
    def stop_workflow(self):
        """Stop the current workflow"""
        if not hasattr(self, 'worker') or not self.worker:
            return
            
        # Confirm before stopping
        reply = QMessageBox.question(
            self, '작업 중지',
            '작업을 중지하시겠습니까? 진행 중인 작업은 저장되지 않을 수 있습니다.',
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            self.worker.stop()
            self.worker.quit()
            self.worker.wait(2000)  # Wait up to 2 seconds
            if self.worker.isRunning():
                self.worker.terminate()
                self.worker.wait()
            
            self.workflow_running = False
            self.workflow_paused = False
            self.update_ui_state()
            self.ui.logTextEdit.append('작업이 중지되었습니다.')
    
    def update_ui_state(self):
        """Update UI elements based on workflow state"""
        self.ui.runButton.setEnabled(not self.workflow_running)
        self.pauseButton.setEnabled(self.workflow_running)
        self.stopButton.setEnabled(self.workflow_running)
        
        # Disable input fields when workflow is running
        for widget in [
            self.ui.appNameLineEdit, self.ui.windowLabelLineEdit,
            self.ui.outputDirLineEdit, self.ui.bookLineEdit,
            self.ui.startSpinBox, self.ui.noSpinBox,
            self.ui.delaySpinBox, self.ui.widthSpinBox, 
            self.ui.heightSpinBox, self.ui.topSpinBox, 
            self.ui.bottomSpinBox, self.ui.leftSpinBox, 
            self.ui.rightSpinBox, self.ui.tessPathLineEdit, 
            self.ui.langLineEdit, self.ui.ocrCheckBox, 
            self.ui.pdfCheckBox, self.ui.textCheckBox, 
            self.ui.mergeCheckBox, self.ui.browseOutputDirButton
        ]:
            widget.setEnabled(not self.workflow_running)
            
    def on_workflow_finished(self):
        """Slot to handle workflow completion"""
        self.workflow_running = False
        self.workflow_paused = False
        self.update_ui_state()
        self.ui.logTextEdit.append('모든 작업이 완료되었습니다.')
        
    def start_workflow(self):
        """
        작업 스레드를 시작하는 메서드
        
        사용자 입력을 검증하고, Worker 스레드를 생성하여 작업을 시작합니다.
        """
        if hasattr(self, 'worker') and self.worker and self.worker.isRunning():
            QMessageBox.warning(self, '경고', '이미 작업이 실행 중입니다.')
            return
        
        # Validate start and end page values
        start_page = self.ui.startSpinBox.value()
        end_page = self.ui.noSpinBox.value()
        
        if end_page < start_page:
            QMessageBox.warning(
                self, 
                '경고', 
                f'시작 페이지({start_page})가 종료 페이지({end_page})보다 큽니다.\n' \
                '종료 페이지를 시작 페이지보다 크거나 같게 설정해주세요.'
            )
            self.ui.noSpinBox.setValue(start_page)
            self.ui.noSpinBox.setFocus()
            return
            
        if start_page < 1 or end_page < 1:
            QMessageBox.warning(
                self,
                '경고',
                '페이지 번호는 1 이상이어야 합니다.'
            )
            return
        
        self.workflow_running = True
        self.workflow_paused = False
        self.update_ui_state()
        
        # Add log message with capture range
        self.ui.logTextEdit.append(f'캡처를 시작합니다: 페이지 {start_page} ~ {end_page}')
        
        try:
            params = self.get_params()
            self.worker = Worker(params)
            self.worker.log_signal.connect(self.ui.logTextEdit.append)
            self.worker.error_signal.connect(lambda msg: QMessageBox.critical(self, '오류', msg))
            self.worker.finished_signal.connect(self.on_workflow_finished)
            self.worker.start()
        except Exception as e:
            self.workflow_running = False
            self.update_ui_state()
            QMessageBox.critical(self, '오류', f'작업을 시작하는 중 오류가 발생했습니다:\n{str(e)}')

if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    window = DocumentProcessorGUI()
    window.show()
    sys.exit(app.exec())